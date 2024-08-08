import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Report:
    def __init__(self, type, params_input, params_output):
        self.type = type
        self.params_input = params_input
        self.params_output = params_output
        self.Document = Document()

    def add_paragraph(self, text):
        self.Document.add_paragraph(text)

    def add_bullets(self, items):
        for item in items:
            self.Document.add_paragraph(item, style='List Bullet')

    def add_header(self, text):
        section = self.Document.sections[0]
        header = section.header

        # Add a paragraph to the header with your desired content
        paragraph = header.paragraphs[0]
        paragraph.text = text

    def add_footer(self, text_1, text_2):
        section = self.Document.sections[0]
        footer = section.footer

        # Add a paragraph to the footer with your desired content
        paragraph = footer.paragraphs[0]
        paragraph.text = text_1
        #paragraph.text = f"{text_1}\tCenter Text\t{text_2}"
        paragraph.text = f"{text_1}\t\t{text_2}"

        # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    def add_table(self, cols, dict):
        table = self.Document.add_table(rows=1, cols=cols)
        first_row = table.rows[0].cells
        first_row[0].text = 'Field'
        first_row[1].text = 'Value'
        table.style = 'Table Grid'

        for col, value in dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(col)
            row_cells[1].text = str(round(value,2))

    def add_image(self, image_path):
        self.Document.add_picture(image_path, width=Inches(4))

    def export_report(self, full_name):
        self.Document.save(full_name)





